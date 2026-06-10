import streamlit as st
import shutil
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import os
import re

st.set_page_config(page_title="Siemens Offer Letter Generator", page_icon="assets/siemens_icon.png", layout="wide")

# ─────────────────────────────────────────────────────────────
# SIEMENS PROFESSIONAL THEME
# ─────────────────────────────────────────────────────────────
st.markdown("""
&lt;style&gt;
    /* ── Google Font: Source Sans 3 (clean, corporate) ── */
    @import url('https://fonts.googleapis.com/css2?family=Source+Sans+3:wght@300;400;500;600;700&amp;display=swap');

    /* ── Global reset ── */
    html, body, [class*="css"] {
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }

    /* ── Page background ── */
    .stApp {
        background-color: #f0f2f5;
    }

    /* ── Hide Streamlit default chrome ── */
    #MainMenu, footer, header { visibility: hidden; }

    /* ── Top banner ── */
    .siemens-banner {
        background: linear-gradient(135deg, #009999 0%, #007a7a 60%, #005f5f 100%);
        border-radius: 6px;
        padding: 26px 36px;
        margin-bottom: 24px;
        display: flex;
        align-items: center;
        justify-content: space-between;
        box-shadow: 0 4px 18px rgba(0,153,153,0.22);
    }
    .siemens-banner-left h1 {
        color: #ffffff;
        font-size: 1.75rem;
        font-weight: 700;
        margin: 0;
        letter-spacing: 0.2px;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }
    .siemens-banner-left p {
        color: rgba(255,255,255,0.78);
        font-size: 0.85rem;
        margin: 5px 0 0 0;
        font-weight: 400;
        letter-spacing: 0.3px;
    }
    .siemens-logo-box {
        background: rgba(255,255,255,0.12);
        border-radius: 4px;
        padding: 8px 22px;
        color: #ffffff;
        font-size: 1.35rem;
        font-weight: 800;
        letter-spacing: 3px;
        border: 1px solid rgba(255,255,255,0.28);
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }

    /* ── Section cards ── */
    .section-card {
        background: #ffffff;
        border-radius: 6px;
        padding: 24px 28px 20px 28px;
        margin-bottom: 18px;
        border: 1px solid #dde2ea;
        box-shadow: 0 1px 5px rgba(0,0,0,0.05);
    }

    /* ── Section header ── */
    .section-header {
        font-size: 0.78rem;
        font-weight: 700;
        color: #009999;
        text-transform: uppercase;
        letter-spacing: 1.2px;
        margin-bottom: 18px;
        padding-bottom: 10px;
        border-bottom: 2px solid #009999;
        display: flex;
        align-items: center;
        gap: 10px;
    }

    /* ── Divider ── */
    hr { border: none; border-top: 1px solid #e4e8ef; margin: 18px 0; }

    /* ── Input labels ── */
    label, .stSelectbox label, .stTextInput label,
    .stNumberInput label, .stTextArea label, .stRadio label {
        font-size: 0.78rem !important;
        font-weight: 600 !important;
        color: #4a5568 !important;
        text-transform: uppercase !important;
        letter-spacing: 0.5px !important;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif !important;
    }

    /* ── Input boxes ── */
    .stTextInput &gt; div &gt; div &gt; input,
    .stTextArea &gt; div &gt; div &gt; textarea,
    .stNumberInput &gt; div &gt; div &gt; input {
        border: 1.5px solid #cbd5e0 !important;
        border-radius: 4px !important;
        font-size: 0.9rem !important;
        color: #1a202c !important;
        background: #fafbfc !important;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif !important;
        transition: border-color 0.15s, box-shadow 0.15s;
    }
    .stTextInput &gt; div &gt; div &gt; input:focus,
    .stTextArea &gt; div &gt; div &gt; textarea:focus {
        border-color: #009999 !important;
        background: #ffffff !important;
        box-shadow: 0 0 0 3px rgba(0,153,153,0.10) !important;
    }

    /* ── Selectbox ── */
    .stSelectbox &gt; div &gt; div {
        border: 1.5px solid #cbd5e0 !important;
        border-radius: 4px !important;
        background: #fafbfc !important;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif !important;
    }

    /* ── Radio buttons ── */
    .stRadio &gt; div {
        gap: 8px;
    }
    .stRadio &gt; div &gt; label {
        background: #f7f9fc;
        border: 1.5px solid #dde2ea;
        border-radius: 4px;
        padding: 8px 18px !important;
        font-size: 0.88rem !important;
        font-weight: 500 !important;
        text-transform: none !important;
        letter-spacing: 0 !important;
        color: #2d3748 !important;
        transition: all 0.15s;
        cursor: pointer;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif !important;
    }
    .stRadio &gt; div &gt; label:hover {
        border-color: #009999;
        background: #f0fafa;
        color: #007a7a !important;
    }

    /* ── Primary button ── */
    .stButton &gt; button[kind="primary"] {
        background: linear-gradient(135deg, #009999 0%, #007a7a 100%) !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 4px !important;
        font-size: 0.95rem !important;
        font-weight: 700 !important;
        padding: 14px 0 !important;
        letter-spacing: 0.8px;
        text-transform: uppercase;
        box-shadow: 0 3px 12px rgba(0,153,153,0.30) !important;
        transition: all 0.15s !important;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif !important;
    }
    .stButton &gt; button[kind="primary"]:hover {
        background: linear-gradient(135deg, #00b3b3 0%, #009999 100%) !important;
        box-shadow: 0 5px 18px rgba(0,153,153,0.40) !important;
        transform: translateY(-1px);
    }

    /* ── Download button ── */
    .stDownloadButton &gt; button {
        background: #ffffff !important;
        color: #009999 !important;
        border: 2px solid #009999 !important;
        border-radius: 4px !important;
        font-size: 0.92rem !important;
        font-weight: 700 !important;
        padding: 12px 0 !important;
        letter-spacing: 0.5px;
        text-transform: uppercase;
        transition: all 0.15s !important;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif !important;
    }
    .stDownloadButton &gt; button:hover {
        background: #009999 !important;
        color: #ffffff !important;
        box-shadow: 0 3px 12px rgba(0,153,153,0.28) !important;
    }

    /* ── Alert / info boxes ── */
    .stAlert[data-baseweb="notification"] {
        border-radius: 4px !important;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif !important;
    }

    /* ── Number input ── */
    .stNumberInput &gt; div { border-radius: 4px !important; }

    /* ── Caption / helper text ── */
    .stCaption, small {
        color: #718096 !important;
        font-size: 0.80rem !important;
    }

    /* ── Scope item cards ── */
    .scope-item-card {
        background: #f7f9fc;
        border: 1px solid #dde2ea;
        border-left: 3px solid #009999;
        border-radius: 4px;
        padding: 14px 18px;
        margin-bottom: 12px;
    }
    .scope-item-label {
        font-size: 0.78rem;
        font-weight: 700;
        color: #4a5568;
        text-transform: uppercase;
        letter-spacing: 0.8px;
        margin-bottom: 10px;
    }

    /* ── Filename preview pill ── */
    .filename-pill {
        display: inline-block;
        background: #e6f7f7;
        border: 1px solid #009999;
        border-radius: 3px;
        padding: 5px 16px;
        font-size: 0.82rem;
        font-weight: 600;
        color: #007a7a;
        margin-top: 8px;
        letter-spacing: 0.2px;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }

    /* ── Contact info boxes ── */
    .contact-info-box {
        background: #f0fafa;
        border: 1px solid #009999;
        border-radius: 4px;
        padding: 10px 14px;
        font-size: 0.87rem;
        color: #1a202c;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }
    .contact-info-box span {
        font-weight: 600;
        color: #007a7a;
        font-size: 0.75rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        display: block;
        margin-bottom: 2px;
    }

    /* ── Footer ── */
    .siemens-footer {
        text-align: center;
        color: #a0aec0;
        font-size: 0.75rem;
        padding: 22px 0 8px 0;
        border-top: 1px solid #dde2ea;
        margin-top: 30px;
        letter-spacing: 0.3px;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }
    .siemens-footer strong {
        color: #009999;
        font-weight: 600;
    }

    /* ── Step badge ── */
    .step-badge {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 22px;
        height: 22px;
        background: #009999;
        color: white;
        border-radius: 50%;
        font-size: 0.70rem;
        font-weight: 700;
        flex-shrink: 0;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }

    /* ── Offer type info box ── */
    .offer-type-info {
        background: #f0fafa;
        border-left: 3px solid #009999;
        border-radius: 0 4px 4px 0;
        padding: 10px 16px;
        font-size: 0.88rem;
        color: #2d3748;
        margin-top: 10px;
        font-family: 'Source Sans 3', 'Segoe UI', Arial, sans-serif;
    }
&lt;/style&gt;
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# TOP BANNER
# ─────────────────────────────────────────────────────────────
st.markdown("""
&lt;div class="siemens-banner"&gt;
    &lt;div class="siemens-banner-left"&gt;
        &lt;h1&gt;Offer Letter Generator&lt;/h1&gt;
        &lt;p&gt;Siemens Industrial LLC &amp;nbsp;&amp;middot;&amp;nbsp; Smart Document Automation &amp;nbsp;&amp;middot;&amp;nbsp; RC-AE SI EA S&lt;/p&gt;
    &lt;/div&gt;
    &lt;div class="siemens-logo-box"&gt;SIEMENS&lt;/div&gt;
&lt;/div&gt;
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# BASE DIRECTORY
# ─────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────

def number_to_words(n):
    ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
            "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
            "Seventeen", "Eighteen", "Nineteen"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]

    def helper(num):
        if num == 0:
            return ""
        elif num &lt; 20:
            return ones[num]
        elif num &lt; 100:
            return tens[num // 10] + (" " + ones[num % 10] if num % 10 != 0 else "")
        elif num &lt; 1000:
            return ones[num // 100] + " Hundred" + (" " + helper(num % 100) if num % 100 != 0 else "")
        elif num &lt; 1_000_000:
            return helper(num // 1000) + " Thousand" + (" " + helper(num % 1000) if num % 1000 != 0 else "")
        elif num &lt; 1_000_000_000:
            return helper(num // 1_000_000) + " Million" + (" " + helper(num % 1_000_000) if num % 1_000_000 != 0 else "")
        else:
            return helper(num // 1_000_000_000) + " Billion" + (" " + helper(num % 1_000_000_000) if num % 1_000_000_000 != 0 else "")

    try:
        clean = str(n).replace(",", "").replace(" ", "").strip()
        num = int(float(clean))
        if num == 0:
            return "Zero"
        return helper(num)
    except:
        return ""

def clear_highlight(run):
    rPr = run._r.get_or_add_rPr()
    for hl in rPr.findall(qn("w:highlight")):
        rPr.remove(hl)

def fill(run, text):
    run.text = text
    run.font.highlight_color = None
    clear_highlight(run)

def all_paras(doc):
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(elem, doc)
        elif tag == "tbl":
            for row in Table(elem, doc).rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr is None:
                continue
            for p in hdr.paragraphs:
                yield p
            for tbl in hdr.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

def merge_and_replace(para, replacements):
    runs = para.runs
    if not runs:
        return

    changed = True
    while changed:
        changed = False
        runs = para.runs
        for i in range(len(runs) - 1):
            r0, r1 = runs[i], runs[i + 1]
            if r0.font.highlight_color is not None and r1.font.highlight_color is not None:
                combined = (r0.text or "") + (r1.text or "")
                if combined.strip() in replacements or "INSERT_" in combined:
                    r0.text = combined
                    r1.text = ""
                    clear_highlight(r1)
                    r1.font.highlight_color = None
                    changed = True
                    break

    for r in para.runs:
        if not r.text:
            continue
        key = r.text.strip()
        if r.font.highlight_color is not None:
            if key in replacements:
                fill(r, replacements[key])
            elif "INSERT_" in r.text:
                new_text = re.sub(
                    r"INSERT_\w+",
                    lambda m: replacements.get(m.group(0), ""),
                    r.text
                )
                fill(r, new_text)

def fill_table(table, items):
    for ri, item in enumerate(items):
        while ri + 1 &gt;= len(table.rows):
            table.add_row()
        row = table.rows[ri + 1]
        vals = [item.get("no", str(ri + 1)), item.get("desc", ""),
                item.get("qty", ""), item.get("total", "")]
        for ci, val in enumerate(vals):
            if ci &lt; len(row.cells):
                p = row.cells[ci].paragraphs[0]
                if p.runs:
                    p.runs[0].text = val
                else:
                    r2 = p.add_run(val)
                    r2.font.name = "Arial"
                    r2.font.size = Pt(9)

# ─────────────────────────────────────────────────────────────
# SECTION 1 - OFFER TYPE
# ─────────────────────────────────────────────────────────────

st.markdown('&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;1&lt;/span&gt;Offer Type&lt;/div&gt;', unsafe_allow_html=True)
offer_type = st.radio("Select the type of commercial offer to generate:", ["Firm", "Budgetary"], horizontal=True)
is_firm = (offer_type == "Firm")
if is_firm:
    st.markdown('&lt;div class="offer-type-info"&gt;&lt;strong&gt;Firm Offer&lt;/strong&gt; — Legally binding commercial offer with full terms, MFC date, validity, and signatories.&lt;/div&gt;', unsafe_allow_html=True)
else:
    st.markdown('&lt;div class="offer-type-info"&gt;&lt;strong&gt;Budgetary Offer&lt;/strong&gt; — Non-binding indicative price offer for budgeting purposes only.&lt;/div&gt;', unsafe_allow_html=True)
st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# SECTION 2 - CUSTOMER INFORMATION
# ─────────────────────────────────────────────────────────────

st.markdown('&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;2&lt;/span&gt;Customer Information&lt;/div&gt;', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    company_options = ["Electro Mechanical Co. LLC (ELMEC)", "ADNOC", "Nofal for Trade & Agencies", "Siemens Energy LLC", "Other"]
    customer_company_sel = st.selectbox("Customer Company", company_options)
    customer_company = st.text_input("Enter Company Name", placeholder="Type company name...") if customer_company_sel == "Other" else customer_company_sel

with col2:
    country_options = ["UAE", "Saudi Arabia", "Algeria", "Yemen", "Kuwait", "Oman", "Qatar", "Other"]
    country_sel = st.selectbox("Country", country_options)
    country = st.text_input("Enter Country Name", placeholder="Type country name...") if country_sel == "Other" else country_sel

col3, col4 = st.columns(2)
with col3:
    customer_attn = st.text_input("Contact Person", placeholder="e.g. Mr. John Smith")
with col4:
    city_default = "Abu Dhabi" if country in ["UAE", ""] else country
    customer_city_input = st.text_input("City", value=city_default)

col5, col6 = st.columns(2)
with col5:
    customer_po_box = st.text_input("P.O. Box", placeholder="Leave blank if not applicable")
with col6:
    customer_fax = st.text_input("Fax", placeholder="Leave blank if not applicable")

col7, col8 = st.columns(2)
with col7:
    customer_tel = st.text_input("Tel", placeholder="Leave blank if not applicable")
with col8:
    customer_mob = st.text_input("Mobile", placeholder="Leave blank if not applicable")

col9, col10 = st.columns(2)
with col9:
    reference_no = st.text_input("Reference No.", placeholder="Leave blank if not applicable")
with col10:
    offer_date = st.text_input("Offer Date", placeholder="e.g. 15 April 2025")

st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# SECTION 3 - OFFER DETAILS
# ─────────────────────────────────────────────────────────────

st.markdown('&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;3&lt;/span&gt;Offer Details&lt;/div&gt;', unsafe_allow_html=True)
subject = st.text_input("Subject", placeholder="e.g. 33KV Switchgear Supply")
col_od1, col_od2 = st.columns(2)
with col_od1:
    project_name = st.text_input("Project Name", placeholder="e.g. PDHPP Project")
with col_od2:
    end_user = st.text_input("End User", placeholder="e.g. STEP Spa")
st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# SECTION 4 - COMMERCIAL TERMS
# ─────────────────────────────────────────────────────────────

st.markdown('&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;4&lt;/span&gt;Commercial Terms&lt;/div&gt;', unsafe_allow_html=True)

col_c1, col_c2 = st.columns(2)
with col_c1:
    currency_options = {
        "EUR - Euro": ("EUR", "Euro"),
        "USD - US Dollar": ("USD", "US Dollar"),
        "GBP - British Pound": ("GBP", "British Pound"),
        "AED - UAE Dirham": ("AED", "UAE Dirham"),
        "SAR - Saudi Riyal": ("SAR", "Saudi Riyal"),
    }
    currency_sel = st.selectbox("Currency", list(currency_options.keys()))
    currency_code, currency_full = currency_options[currency_sel]
with col_c2:
    total_price_num = st.text_input("Total Price", placeholder="e.g. 3,218,545")

col_c3, col_c4 = st.columns(2)
with col_c3:
    incoterm_options = ["FCA, Bremerhaven, Germany", "CIF, Abu Dhabi seaport",
                        "DAP, Abu Dhabi", "EXW, Bremerhaven, Germany", "Other"]
    incoterm_sel = st.selectbox("Incoterm", incoterm_options)
    incoterm_name = st.text_input("Enter Incoterm", placeholder="e.g. FCA, Hamburg, Germany") if incoterm_sel == "Other" else incoterm_sel
with col_c4:
    delivery_options = ["06", "09", "12", "15", "18", "24", "Other"]
    delivery_sel = st.selectbox("Delivery Period (Months)", delivery_options)
    delivery_months = st.text_input("Enter Months", placeholder="e.g. 10") if delivery_sel == "Other" else delivery_sel

col_c5, col_c6 = st.columns(2)
with col_c5:
    warranty_options = ["12", "18", "24", "Other"]
    warranty_sel = st.selectbox("Warranty Period (Months)", warranty_options)
    warranty_months = st.text_input("Enter Warranty Months", placeholder="e.g. 24") if warranty_sel == "Other" else warranty_sel

st.markdown("---")
st.markdown("**Payment Structure**")

payment_option_labels = [
    "Option A — 20% down payment + 80% against shipping documents",
    "Option B — 20% advance + 10% drawings + 20% MFC + 50% bill of lading",
    "Option C — Custom payment terms",
]
payment_option_sel = st.radio(
    "Select payment structure",
    payment_option_labels,
    label_visibility="collapsed",
)

if payment_option_sel.startswith("Option A"):
    payment_option = "A"
    custom_payment_text = ""
    col_pd1, col_pd2 = st.columns([1, 3])
    with col_pd1:
        payment_days_options = ["30", "45", "60", "90", "120", "Other"]
        payment_days_sel = st.selectbox("Payment Days", payment_days_options, key="pd_A")
        payment_days = st.text_input("Enter Days", key="pd_A_custom") if payment_days_sel == "Other" else payment_days_sel

elif payment_option_sel.startswith("Option B"):
    payment_option = "B"
    custom_payment_text = ""
    col_pd1, col_pd2 = st.columns([1, 3])
    with col_pd1:
        payment_days_options = ["30", "45", "60", "90", "120", "Other"]
        payment_days_sel = st.selectbox("Payment Days", payment_days_options, key="pd_B")
        payment_days = st.text_input("Enter Days", key="pd_B_custom") if payment_days_sel == "Other" else payment_days_sel

else:
    payment_option = "C"
    custom_payment_text = st.text_area(
        "Enter Custom Payment Terms",
        placeholder=(
            "e.g.\n"
            "100% of the contract value payable against the date of "
            "delivery documents within 90 days."
        ),
        height=120,
        key="custom_payment_text",
    )
    col_pd1, col_pd2 = st.columns([1, 3])
    with col_pd1:
        payment_days_options = ["30", "45", "60", "90", "120", "Other"]
        payment_days_sel = st.selectbox(
            "Payment Days (for Header)",
            payment_days_options,
            key="pd_C",
            help="Used in the section header line only",
        )
        payment_days = (
            st.text_input("Enter Days", key="pd_C_custom")
            if payment_days_sel == "Other"
            else payment_days_sel
        )
    with col_pd2:
        st.info(
            f"Header will read: **Payment terms: {payment_days} days from delivery documents**\n\n"
            "The text above will be inserted verbatim into the document."
        )

if is_firm:
    st.markdown("---")
    cancel_high = st.selectbox(
        "Cancellation Charge — High Bracket",
        ["80%", "90%"],
        help="Select -80% for National customers, -90% for International / Siemens Energy / Critical Country"
    )
else:
    cancel_high = ""

st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# SECTION 5 - FIRM ONLY FIELDS
# ─────────────────────────────────────────────────────────────

if is_firm:
    st.markdown('&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;5&lt;/span&gt;Firm Offer Details&lt;/div&gt;', unsafe_allow_html=True)

    install_options = {
        "UAE - Abu Dhabi (Dubai or Northern Emirates)": "the Emirate of Abu Dhabi (Dubai or Northern Emirates)",
        "UAE - Dubai / Jebel Ali": "the Emirate of Dubai / Jebel Ali",
        "Saudi Arabia": "a Saudi Arabian seaport",
        "Algeria": "an Algerian seaport",
        "Yemen": "a Yemeni seaport",
        "Other": None
    }
    install_sel = st.selectbox("Country / Port of Installation", list(install_options.keys()))
    import_port = st.text_input("Enter Country / Port", placeholder="e.g. a Kuwaiti seaport") if install_sel == "Other" else install_options[install_sel]

    col_f1, col_f2 = st.columns(2)
    with col_f1:
        mfc_date = st.text_input("MFC Date", placeholder="e.g. 26th February 2026")
    with col_f2:
        offer_validity = st.text_input("Offer Validity", placeholder="e.g. March 30, 2026")

    signatory_options = [
        "Robert Hennig + Rodrigo Fernandes",
        "Robert Hennig + Rodrigo Fernandes + Prashanth Parameswaran + Paul Fairweather",
        "Robert Hennig + Rodrigo Fernandes + Prashanth Parameswaran + Paul Fairweather + Nasir Merchant",
        "Other"
    ]
    signatory_sel = st.selectbox("Signatories", signatory_options)
    signatories = st.text_input("Enter Signatory Names", placeholder="Name 1 + Name 2 + ...") if signatory_sel == "Other" else signatory_sel
    st.markdown('&lt;/div&gt;', unsafe_allow_html=True)
else:
    import_port    = ""
    mfc_date       = ""
    offer_validity = ""
    signatories    = ""

# ─────────────────────────────────────────────────────────────
# SECTION 6 - SALES CONTACT
# ─────────────────────────────────────────────────────────────

step_num = "6" if is_firm else "5"
st.markdown(f'&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;{step_num}&lt;/span&gt;Sales Contact&lt;/div&gt;', unsafe_allow_html=True)

sales_contacts = {
    "Ahmad Awny | RC-AE SI EA S VD-V-D | +971 55 2003541 | ahmad.awny@siemens.com":
        {"name": "Ahmad Awny", "dept": "RC-AE SI EA S VD-V-D", "mobile": "+971 55 2003541", "email": "ahmad.awny@siemens.com"},
    "Vivek Gopalakrishnan | RC-AE SI EA S | +971 55 2002583 | vivek.gopalakrishnan@siemens.com":
        {"name": "Vivek Gopalakrishnan", "dept": "RC-AE SI EA S", "mobile": "+971 55 2002583", "email": "vivek.gopalakrishnan@siemens.com"},
    "Madiha Khan | RC-AE SI EA S | +971 55 2003818 | madiha.khan@siemens.com":
        {"name": "Madiha Khan", "dept": "RC-AE SI EA S", "mobile": "+971 55 2003818", "email": "madiha.khan@siemens.com"},
    "Hyun-Sik Kim | RC-AE SI EA S VD-V-D | +971 55 2002693 | hyun-sik.kim@siemens.com":
        {"name": "Hyun-Sik Kim", "dept": "RC-AE SI EA S VD-V-D", "mobile": "+971 55 2002693", "email": "hyun-sik.kim@siemens.com"},
    "Other": None
}

sales_sel = st.selectbox("Sales Contact", list(sales_contacts.keys()))
if sales_sel == "Other":
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        sender_name   = st.text_input("Sender Name", placeholder="Full name")
        sender_mobile = st.text_input("Sender Mobile", placeholder="+971 XX XXXXXXX")
    with col_s2:
        sender_dept  = st.text_input("Sender Department", placeholder="e.g. RC-AE SI EA S")
        sender_email = st.text_input("Sender Email", placeholder="name@siemens.com")
else:
    sc = sales_contacts[sales_sel]
    sender_name   = sc["name"]
    sender_dept   = sc["dept"]
    sender_mobile = sc["mobile"]
    sender_email  = sc["email"]
    col_si1, col_si2, col_si3 = st.columns(3)
    with col_si1:
        st.markdown(f"&lt;div class='contact-info-box'&gt;&lt;span&gt;Email&lt;/span&gt;{sender_email}&lt;/div&gt;", unsafe_allow_html=True)
    with col_si2:
        st.markdown(f"&lt;div class='contact-info-box'&gt;&lt;span&gt;Mobile&lt;/span&gt;{sender_mobile}&lt;/div&gt;", unsafe_allow_html=True)
    with col_si3:
        st.markdown(f"&lt;div class='contact-info-box'&gt;&lt;span&gt;Department&lt;/span&gt;{sender_dept}&lt;/div&gt;", unsafe_allow_html=True)

st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# SECTION 7 - SCOPE OF SUPPLY
# ─────────────────────────────────────────────────────────────

step_num2 = "7" if is_firm else "6"
st.markdown(f'&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;{step_num2}&lt;/span&gt;Scope of Supply&lt;/div&gt;', unsafe_allow_html=True)

num_scope = st.number_input("Number of Scope Items", min_value=1, max_value=20, value=1, step=1)
scope_items = []
for i in range(int(num_scope)):
    st.markdown(f'&lt;div class="scope-item-card"&gt;&lt;div class="scope-item-label"&gt;Item {i+1}&lt;/div&gt;', unsafe_allow_html=True)
    col_sc1, col_sc2, col_sc3 = st.columns([4, 1, 2])
    with col_sc1: desc  = st.text_input("Description", key=f"scope_desc_{i}", placeholder="e.g. 33KV 8DA10 AIS NXAIR Switchgear Panel")
    with col_sc2: qty   = st.text_input("Qty",         key=f"scope_qty_{i}",  placeholder="e.g. 10")
    with col_sc3: total = st.text_input("Total Price", key=f"scope_total_{i}", placeholder="e.g. EUR 3,218,545")
    scope_items.append({"no": str(i+1), "desc": desc, "qty": qty, "total": total})
    st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

st.markdown("---")
st.markdown("**Optional Items**")
num_opt = st.number_input("Number of Optional Items", min_value=0, max_value=10, value=0, step=1)
optional_items = []
for i in range(int(num_opt)):
    st.markdown(f'&lt;div class="scope-item-card"&gt;&lt;div class="scope-item-label"&gt;Optional Item {i+1}&lt;/div&gt;', unsafe_allow_html=True)
    col_o1, col_o2, col_o3 = st.columns([4, 1, 2])
    with col_o1: odesc  = st.text_input("Description", key=f"opt_desc_{i}")
    with col_o2: oqty   = st.text_input("Qty",         key=f"opt_qty_{i}")
    with col_o3: ototal = st.text_input("Total Price", key=f"opt_total_{i}")
    optional_items.append({"no": str(i+1), "desc": odesc, "qty": oqty, "total": ototal})
    st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# SECTION 8 - OUTPUT FILENAME
# ─────────────────────────────────────────────────────────────

step_num3 = "8" if is_firm else "7"
st.markdown(f'&lt;div class="section-card"&gt;&lt;div class="section-header"&gt;&lt;span class="step-badge"&gt;{step_num3}&lt;/span&gt;Output Filename&lt;/div&gt;', unsafe_allow_html=True)

offer_short_prev = "FIRM" if is_firm else "BUD"
cust_short_prev  = customer_company.split()[0].upper().strip(".,)") if customer_company else "CUSTOMER"
SKIP_WORDS       = {"FOR", "THE", "OF", "AND", "IN", "A", "AN", "PKG", "WORKS", "-", "PROJECT"}
proj_short_prev  = "".join(w for w in project_name.split() if w.upper() not in SKIP_WORDS)[:12].upper() if project_name else "PROJECT"
try:    date_str_prev = datetime.strptime(offer_date.strip(), "%d %B %Y").strftime("%Y%m%d")
except: date_str_prev = offer_date.replace(" ", "")[:8] if offer_date else "DATE"

auto_filename = f"{offer_short_prev}_{cust_short_prev}_{proj_short_prev}_{date_str_prev}_v01"

filename_mode = st.radio(
    "Filename mode",
    ["Auto", "Custom"],
    horizontal=True,
    label_visibility="collapsed",
)

if filename_mode == "Custom":
    custom_filename_input = st.text_input(
        "Enter Filename (without extension)",
        placeholder=f"e.g. {auto_filename}",
    )
    raw = custom_filename_input.strip().removesuffix(".docx")
    final_filename = (raw if raw else auto_filename) + ".docx"
else:
    final_filename = auto_filename + ".docx"

st.markdown(f'&lt;div class="filename-pill"&gt;Output file: {final_filename}&lt;/div&gt;', unsafe_allow_html=True)
st.markdown('&lt;/div&gt;', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# GENERATE
# ─────────────────────────────────────────────────────────────

st.markdown("&lt;br&gt;", unsafe_allow_html=True)
if st.button("Generate Offer Letter", type="primary", use_container_width=True):

    errors = []
    if not customer_company or customer_company == "Other": errors.append("Customer Company")
    if not customer_attn:   errors.append("Contact Person")
    if not offer_date:      errors.append("Offer Date")
    if not subject:         errors.append("Subject")
    if not project_name:    errors.append("Project Name")
    if not end_user:        errors.append("End User")
    if not total_price_num: errors.append("Total Price")
    if payment_option == "C" and not custom_payment_text.strip():
        errors.append("Custom Payment Terms (Option C cannot be empty)")
    if is_firm and not mfc_date:       errors.append("MFC Date")
    if is_firm and not offer_validity: errors.append("Offer Validity")
    if errors:
        st.error(f"Please fill in the following required fields: {', '.join(errors)}")
        st.stop()

    customer_city     = f"{customer_city_input}, {country}"
    total_price_words = number_to_words(total_price_num)

    if payment_option == "A":
        pay_header = f"Payment terms: {payment_days} days from shipping documents"
        pay_lines  = (
            f"20% of the contract value to be paid as down payment within 15 days "
            f"after placement of the order.\n"
            f"80% of the contract value payable against presentation of shipping "
            f"documents or warehouse receipt within {payment_days} days."
        )
    elif payment_option == "B":
        pay_header = f"Payment terms: {payment_days} days from shipping documents"
        pay_lines  = (
            f"20% of the contract value to be paid as Advance Payment within 15 days "
            f"after Order Confirmation.\n"
            f"10% of the contract value against submission of Drawings (SLD - Single "
            f"Line Diagrams), payable within {payment_days} days.\n"
            f"20% of the contract value upon Manufacturing Clearance, payable within "
            f"45 days.\n"
            f"50% of the contract value against presentation of shipping documents "
            f"(Bill of Lading), payable within {payment_days} days."
        )
    else:
        pay_header = f"Payment terms: {payment_days} days from delivery documents"
        pay_lines  = custom_payment_text.strip()

    import_port_sentence = (
        f"The scope shall be offloaded in and imported via a port of {import_port}."
        if is_firm and import_port else ""
    )

    po_box_full = f"P. O. Box {customer_po_box}" if customer_po_box.strip() else ""
    tel_str     = f"Tel : {customer_tel}"         if customer_tel.strip()    else ""
    fax_str     = f"Fax: {customer_fax}"          if customer_fax.strip()    else ""
    mob_str     = f"Mob: {customer_mob}"          if customer_mob.strip()    else ""

    price_words_run = f"{currency_code} {total_price_words} Only)."

    replacements = {
        "INSERT_CUSTOMER_COMPANY":                          customer_company,
        "P. O. Box INSERT_CUSTOMER_PO_BOX_NUM":            po_box_full,
        "INSERT_CUSTOMER_CITY_FULL":                        customer_city,
        "INSERT_CUSTOMER_ATTN":                             f"Kind Attn: {customer_attn}",
        "INSERT_CUSTOMER_TEL_LINE":                         tel_str,
        "INSERT_CUSTOMER_FAX_LINE":                         fax_str,
        "INSERT_CUSTOMER_MOB_LINE":                         mob_str,
        "INSERT_SENDER_NAME":                               sender_name,
        "INSERT_SENDER_DEPT":                               sender_dept,
        "INSERT_SENDER_MOBILE":                             sender_mobile,
        "INSERT_SENDER_EMAIL":                              sender_email,
        "INSERT_OFFER_DATE":                                offer_date,
        "INSERT_REFERENCE_NO":                              reference_no.strip(),
        "INSERT_SUBJECT_FULL":                              subject,
        "INSERT_PROJECT_NAME":                              project_name,
        "INSERT_END_USER":                                  end_user,
        "INSERT_CURRENCY_FULL":                             currency_full,
        "INSERT_TOTAL_PRICE_NUM":                           total_price_num,
        "INSERT_CURRENCY_CODE":                             currency_code,
        "INSERT_CURRENCY_CODE INSERT_TOTAL_PRICE_WORDS).":  price_words_run,
        "INSERT_CURRENCY_CODE INSERT_TOTAL_PRICE_WORDS Only).": price_words_run,
        "INSERT_INCOTERM_NAME":                             incoterm_name,
        "INSERT_DELIVERY_MONTHS":                           f"{delivery_months} month(s)",
        "INSERT_WARRANTY_MONTHS":                           f"{warranty_months} months",
        "INSERT_PAYMENT_OPTION_HEADER":                     pay_header,
        "INSERT_PAYMENT_OPTION_LINE":                       pay_lines,
        "INSERT_PAYMENT_OPTION_LINES":                      pay_lines,
        "INSERT_MFC_DATE":                                  mfc_date if is_firm else "",
        "INSERT_IMPORT_PORT_SENTENCE":                      import_port_sentence,
        "INSERT_OFFER_VALIDITY":                            offer_validity if is_firm else "",
        "INSERT_CANCEL_HIGH":                               cancel_high,
    }

    TEMPLATE_PATH = (
        os.path.join(BASE_DIR, "template_Firm_FIXED.docx")
        if is_firm else
        os.path.join(BASE_DIR, "template_Budgetary_FIXED.docx")
    )

    filepath = os.path.join(BASE_DIR, final_filename)
    shutil.copy(TEMPLATE_PATH, filepath)
    doc = Document(filepath)

    for para in all_paras(doc):
        merge_and_replace(para, replacements)

    for para in all_paras(doc):
        runs = para.runs
        i = 0
        while i &lt; len(runs) - 1:
            r0, r1 = runs[i], runs[i + 1]
            if "INSERT_" in (r0.text or "") and "INSERT_" in (r1.text or ""):
                combined = (r0.text or "") + (r1.text or "")
                if combined.strip() in replacements:
                    r0.text = combined
                    r1.text = ""
                    continue
            i += 1
        for r in para.runs:
            if not r.text or "INSERT_" not in r.text:
                continue
            key = r.text.strip()
            if key in replacements:
                r.text = replacements[key]
                clear_highlight(r)
            else:
                new_text = re.sub(
                    r"INSERT_\w+",
                    lambda m: replacements.get(m.group(0), ""),
                    r.text
                )
                if new_text != r.text:
                    r.text = new_text
                    clear_highlight(r)

    if len(doc.tables) &gt; 1:
        fill_table(doc.tables[1], scope_items)
    if optional_items and len(doc.tables) &gt; 2:
        fill_table(doc.tables[2], optional_items)

    for para in all_paras(doc):
        for r in para.runs:
            if r.text and "INSERT_" in r.text:
                r.text = re.sub(r"INSERT_\w+", "", r.text)
            if r.font.highlight_color is not None:
                clear_highlight(r)
                r.font.highlight_color = None

    doc.save(filepath)

    st.success("Offer letter generated successfully. Your document is ready for download.")
    with open(filepath, "rb") as f:
        st.download_button(
            label=f"Download {final_filename}",
            data=f,
            file_name=final_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# ─────────────────────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────────────────────
st.markdown("""
&lt;div class="siemens-footer"&gt;
    &lt;strong&gt;Siemens Industrial LLC&lt;/strong&gt; &amp;nbsp;&amp;middot;&amp;nbsp; Offer Letter Generator &amp;nbsp;&amp;middot;&amp;nbsp;
    RC-AE SI EA S &amp;nbsp;&amp;middot;&amp;nbsp; Internal Use Only
&lt;/div&gt;
""", unsafe_allow_html=True)
